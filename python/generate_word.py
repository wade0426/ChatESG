
# pip install python-docx Pillow

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import os
import datetime

def generate_word_document(data):
    document = Document()
    
    for index, (group_key, group_data) in enumerate(data.items()):

        # 添加標題
        title = document.add_paragraph()
        # 設定標題對齊方式為居中
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # 創建run並設置格式
        run = title.add_run(group_data['title'])
        run.bold = True
        # 設置字體大小應該在run上，而不是paragraph上
        run.font.size = Pt(22)
        
        # 添加內文
        # 先創建空段落 用於 run 設定格式
        content = document.add_paragraph() 
        # 添加格式化的文本
        run = content.add_run(group_data['generatedResult'].replace('<br>', '\n'))
        run.font.size = Pt(16)
        
        # 添加圖片
        for chart in group_data['charts']:
            if chart['base64']:
                image_path = chart['base64']
                
                if os.path.exists(image_path):
                    # 創建一個新的段落來包含圖片和相關文本
                    p = document.add_paragraph()
                    p.keep_together = True

                    # 添加圖片標題
                    if chart['imageTitle']:
                        run = p.add_run(chart['imageTitle'] + '\n')
                        run.font.size = Pt(16)
                        run.font.bold = True
                        run.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    # 添加圖片
                    run = p.add_run()  # 在段落中創建一個新的文本運行(run)
                    run.add_picture(image_path, width=Inches(6))  # 在運行中添加圖片，設置寬度為6英寸
                    # 解釋：
                    # 1. p.add_run() 創建一個新的文本運行，這允許我們在段落中添加內容
                    # 2. run.add_picture() 方法用於在運行中添加圖片
                    # 4. width=Inches(6) 設置圖片寬度為6英寸，這有助於保持文檔中圖片大小的一致性
                    
                    # 添加圖片說明
                    if chart['imageDescription']:
                        run = p.add_run('\n' + chart['imageDescription'])
                        run.font.size = Pt(12)
                    
                    # 設置段落對齊方式為居中
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    print(f"警告: 圖片文件 {image_path} 不存在")
        
        # 檢查是否為最後一個元素
        if index < len(data) - 1:
            # 如果不是最後一個元素，才添加換頁符
            document.add_page_break()
    
    # 指定一個新的目錄來保存文件
    save_dir = os.path.join(f"{os.getcwd()}", 'generated_reports')
    os.makedirs(save_dir, exist_ok=True)
    
    timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = os.path.join(save_dir, f'generated_report_{timestamp}.docx')
    
    try:
        document.save(filename)
        print(f"文件已保存為: {filename}")
    except PermissionError:
        print(f"無法保存文件 {filename}。請確保您有寫入權限。")
    except Exception as e:
        print(f"保存文件時發生錯誤: {str(e)}")

    return filename



if __name__ == "__main__":
    # 示例數據
    sample_data = {
        'group1': {
            'title': '長官的話',
            'generatedResult': '中科金控 矢志永續  共創價值 \n',
            'charts': [
                {'base64': 'D:\\NTCUST\\Project\\ESG\\ChatESG\\python\\temp_image\\0.png', 'imageTitle': '1', 'imageDescription': '1', 'url': ''},
                {'base64': 'D:\\NTCUST\\Project\\ESG\\ChatESG\\python\\temp_image\\1.png', 'imageTitle': '2', 'imageDescription': '2', 'url': ''}
            ]
        },
        'group2': {
            'title': '測試',
            'generatedResult': '測試',
            'charts': [
                {'base64': 'D:\\NTCUST\\Project\\ESG\\ChatESG\\python\\temp_image\\2.png', 'imageTitle': '1', 'imageDescription': '2', 'url': ''}
            ]
        }
    }

    # 調用函數生成Word文件
    # generate_word_document(sample_data)
    print("當前路徑:", os.getcwd())
    # 當前路徑 退一層
    # os.chdir("../")
    # print("當前路徑:", os.getcwd())

